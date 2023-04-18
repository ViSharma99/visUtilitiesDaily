import os
import PyPDF2
from docx import Document
import docx

pdf_file_path = '/home/viswanatha/ECS766P_220166250.pdf'
docx_file_path = '/home/viswanatha/example.docx'



with open(pdf_file_path, 'rb') as pdf_file:
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    page_count = len(pdf_reader.pages)
    page_text_list = []

    # Iterate over each page and extract its text
    for page_num in range(page_count):
        page = pdf_reader.pages[page_num]
        page_text = page.extract_text()
        page_text_list.append(page_text)

# Create a new Word document and add the extracted text to it
doc = docx.Document()

for page_text in page_text_list:
    doc.add_paragraph(page_text)

doc.save(docx_file_path)
# Open the PDF file in read-binary mode.
# with open(pdf_file_path, 'rb') as pdf_file:
#     # Create a PDF reader object.
#     pdf_reader = PyPDF2.PdfReader(pdf_file)

#     # Create a new Word document.
#     docx_document = Document()  

#     # Loop through each page in the PDF file.
#     for page_num in range(len(pdf_reader.pages)):
#         # Get the text content of the page.
#         page_text = pdf_reader.pages(page_num).extractText()

#         # Add the page text to the Word document.
#         docx_document.add_paragraph(page_text)

#     # Save the Word document to disk.
#     docx_document.save(docx_file_path)

# Open the resulting Word document.
os.system(f"xdg-open {docx_file_path}")